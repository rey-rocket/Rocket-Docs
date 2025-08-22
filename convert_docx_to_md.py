#!/usr/bin/env python3
"""
Word Document to Markdown Converter

This script converts .docx files to markdown format using python-docx.
"""

import sys
import os
from pathlib import Path
import argparse
import re

try:
    from docx import Document
except ImportError as e:
    print(f"Missing required dependency: {e}")
    print("Please install required packages:")
    print("pip install python-docx")
    sys.exit(1)


def convert_docx_to_markdown(docx_path, output_path=None):
    """
    Convert a Word document to markdown format.
    
    Args:
        docx_path (str): Path to the input .docx file
        output_path (str, optional): Path for the output .md file
    
    Returns:
        str: Path to the created markdown file
    """
    docx_path = Path(docx_path)
    
    if not docx_path.exists():
        raise FileNotFoundError(f"Input file not found: {docx_path}")
    
    if not docx_path.suffix.lower() == '.docx':
        raise ValueError("Input file must have .docx extension")
    
    # Generate output path if not provided
    if output_path is None:
        output_path = docx_path.with_suffix('.md')
    else:
        output_path = Path(output_path)
    
    print(f"Converting {docx_path} to {output_path}")
    
    try:
        doc = Document(docx_path)
        markdown_content = []
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                markdown_content.append("")
                continue
            
            # Handle different paragraph styles
            style_name = paragraph.style.name.lower()
            
            if 'heading 1' in style_name:
                markdown_content.append(f"# {text}")
            elif 'heading 2' in style_name:
                markdown_content.append(f"## {text}")
            elif 'heading 3' in style_name:
                markdown_content.append(f"### {text}")
            elif 'heading 4' in style_name:
                markdown_content.append(f"#### {text}")
            elif 'heading 5' in style_name:
                markdown_content.append(f"##### {text}")
            elif 'heading 6' in style_name:
                markdown_content.append(f"###### {text}")
            else:
                # Handle regular paragraphs with formatting
                formatted_text = format_paragraph_text(paragraph)
                markdown_content.append(formatted_text)
        
        # Process tables
        for table in doc.tables:
            markdown_content.append("")
            table_md = convert_table_to_markdown(table)
            markdown_content.extend(table_md)
            markdown_content.append("")
        
        # Write the markdown content to file
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write("\n".join(markdown_content))
        
        print(f"Successfully converted to: {output_path}")
        return str(output_path)
        
    except Exception as e:
        print(f"Error during conversion: {e}")
        raise


def format_paragraph_text(paragraph):
    """Format a paragraph with bold and italic text."""
    text = ""
    for run in paragraph.runs:
        run_text = run.text
        if run.bold and run.italic:
            text += f"***{run_text}***"
        elif run.bold:
            text += f"**{run_text}**"
        elif run.italic:
            text += f"*{run_text}*"
        else:
            text += run_text
    return text


def convert_table_to_markdown(table):
    """Convert a Word table to markdown format."""
    if not table.rows:
        return []
    
    markdown_rows = []
    
    # Process header row
    if table.rows:
        header_cells = [cell.text.strip() for cell in table.rows[0].cells]
        markdown_rows.append("| " + " | ".join(header_cells) + " |")
        markdown_rows.append("| " + " | ".join(["---"] * len(header_cells)) + " |")
        
        # Process data rows
        for row in table.rows[1:]:
            cells = [cell.text.strip() for cell in row.cells]
            markdown_rows.append("| " + " | ".join(cells) + " |")
    
    return markdown_rows


def main():
    parser = argparse.ArgumentParser(description='Convert Word documents to Markdown')
    parser.add_argument('input', help='Input .docx file path')
    parser.add_argument('-o', '--output', help='Output .md file path (optional)')
    
    args = parser.parse_args()
    
    try:
        convert_docx_to_markdown(args.input, args.output)
    except Exception as e:
        print(f"Conversion failed: {e}")
        sys.exit(1)


if __name__ == "__main__":
    main()